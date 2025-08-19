# -*- coding: utf-8 -*-
# Copyright (c) Huawei Technologies Co., Ltd. 2024. All rights reserved.

import os
import re
import zipfile
from dataclasses import dataclass
from typing import Iterator, Callable

from pathlib import Path
import docx
from loguru import logger
from langchain_core.documents import Document
from langchain_community.document_loaders.base import BaseLoader

from mx_rag.document.loader.base_loader import BaseLoader as mxBaseLoader
from mx_rag.llm import Img2TextLLM
from mx_rag.utils.file_check import SecFileCheck
from mx_rag.utils.common import validate_params, BOOL_TYPE_CHECK_TIP


class DocxLoader(BaseLoader, mxBaseLoader):
    """Loading logic for loading documents from docx."""
    EXTENSION = (".docx",)

    @validate_params(
        vlm=dict(validator=lambda x: isinstance(x, Img2TextLLM) or x is None,
                 message="param must be instance of Img2TextLLM or None"),
        image_inline=dict(validator=lambda x: isinstance(x, bool), message=BOOL_TYPE_CHECK_TIP)
    )
    def __init__(self, file_path: str, vlm: Img2TextLLM = None, image_inline: bool = False):
        """Initialize with filepath and options."""
        super().__init__(file_path)
        self.do_ocr = image_inline
        self.vlm = vlm
        self.table_index = 0

    def lazy_load(self) -> Iterator[Document]:
        """Load documents."""
        self._is_document_valid()
        all_text, img_base64_list, image_summaries = [], [], []
        doc = docx.Document(self.file_path)
        for element in doc.element.body:
            if element.tag.endswith("tbl"):
                table_text = self._handle_table(doc.tables[self.table_index])
                self.table_index += 1
                all_text.append(table_text)
            elif element.tag.endswith("p"):
                if "pic:pic" in str(element.xml) and self.do_ocr:
                    logger.debug("pic:pic set to empty str")
                    pic_texts = ""
                    all_text.extend(pic_texts)
                paragraph = docx.text.paragraph.Paragraph(element, doc)
                para_text = paragraph.text
                all_text.append(para_text)

        one_text = "\n\n".join([t for t in all_text])
        yield Document(page_content=one_text, metadata={"source": os.path.basename(self.file_path), "type": "text"})

        if self.vlm:
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_part = rel.target_part
                    image_data = image_part.blob
                    img_base64, image_summary = self._interpret_image(image_data, self.vlm)
                    img_base64_list.extend([img_base64] if image_summary and img_base64 else [])
                    image_summaries.extend([image_summary] if image_summary and img_base64 else [])

            for img_base64, image_summary in zip(img_base64_list, image_summaries):
                yield Document(page_content=image_summary, metadata={"source": os.path.basename(self.file_path),
                                                                     "image_base64": img_base64, "type": "image"})

    def _handle_table(self, element):
        """docx.oxml.table.CT_Tbl"""

        logger.info(f"handle docx table {self.table_index}")
        rows = list(element.rows)
        headers = [cell.text for cell in rows[0].cells]
        data = [[cell.text.replace("\n", " ") for cell in row.cells] for row in rows[1:]]
        if not data:
            res = "；".join(headers)
        else:
            result = ["，".join([f"{x}: {y}" for x, y in zip(headers, subdata)]) for subdata in data]
            res = "；".join(result)
        return res + "。"

    def _is_document_valid(self):
        SecFileCheck(self.file_path, self.MAX_SIZE).check()
        if not self.file_path.endswith(DocxLoader.EXTENSION):
            raise TypeError(f"type '{Path(self.file_path).suffix}' is not support")
        if self._is_zip_bomb():
            raise ValueError(f"file is a risk of zip bombs")
        doc = docx.Document(self.file_path)
        word_count = 0
        for paragraph in doc.paragraphs:
            word_count += len(paragraph.text)
        if word_count > self.MAX_WORD_NUM:
            raise ValueError(f"too many words {word_count}")
